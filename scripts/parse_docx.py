#!/usr/bin/env python3
"""Parse a .docx requirement document into document-ir.yaml, markdown, and image manifest."""
from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import yaml

try:
    from docx import Document
    from docx.document import Document as DocumentObject
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except Exception:  # pragma: no cover - fallback keeps the parser usable without python-docx.
    Document = None
    DocumentObject = object
    Table = object
    Paragraph = object


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
NS_EXT = {
    "w": NS["w"],
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Requirement .docx file")
    parser.add_argument("--workspace", default="workspace")
    args = parser.parse_args()

    src = Path(args.input)
    ws = Path(args.workspace)
    parsed = ws / "parsed"
    tables_dir = parsed / "tables"
    images_dir = parsed / "images"
    parsed.mkdir(parents=True, exist_ok=True)
    tables_dir.mkdir(parents=True, exist_ok=True)
    images_dir.mkdir(parents=True, exist_ok=True)

    blocks, tables, image_manifest = parse_docx(src, images_dir)
    warnings = detect_parser_warnings(src)
    ir = {
        "document_ir": {
            "document_id": "REQ-DOC-001",
            "filename": src.name,
            "parser_version": "concept-design-docx-v2",
            "parser_warnings": warnings,
            "blocks": blocks,
        }
    }
    write_yaml(parsed / "document-ir.yaml", ir)
    write_markdown(parsed / "document.md", blocks)
    for table in tables:
        write_yaml(tables_dir / f"{table['table_id']}.yaml", {"table": table})
    write_yaml(parsed / "image-manifest.yaml", {"image_manifest": image_manifest})
    print(parsed / "document-ir.yaml")
    return 0


def parse_docx(src: Path, images_dir: Path | None = None) -> tuple[list[dict], list[dict], list[dict]]:
    if Document is not None:
        try:
            return parse_docx_with_python_docx(src, images_dir)
        except Exception:
            return parse_docx_with_xml(src, images_dir)
    return parse_docx_with_xml(src, images_dir)


def parse_docx_with_python_docx(src: Path, images_dir: Path | None = None) -> tuple[list[dict], list[dict], list[dict]]:
    doc = Document(str(src))
    blocks: list[dict] = []
    tables: list[dict] = []
    section_path: list[str] = []
    block_no = 0
    table_no = 0
    order = 0
    for child in iter_doc_blocks(doc):
        if isinstance(child, Paragraph):
            text = re.sub(r"\s+", " ", child.text or "").strip()
            if not text:
                continue
            order += 1
            block_no += 1
            block_id = f"B-{block_no:04d}"
            level = paragraph_heading_level(child)
            if level:
                section_path = section_path[: level - 1] + [text]
                blocks.append(
                    {
                        "block_id": block_id,
                        "block_type": "heading",
                        "level": level,
                        "text": text,
                        "section_path": list(section_path),
                        "order": order,
                    }
                )
            else:
                blocks.append(
                    {
                        "block_id": block_id,
                        "block_type": "paragraph",
                        "text": text,
                        "section_path": list(section_path),
                        "order": order,
                    }
                )
        elif isinstance(child, Table):
            rows = [[re.sub(r"\s+", " ", cell.text or "").strip() for cell in row.cells] for row in child.rows]
            if not any(any(cell for cell in row) for row in rows):
                continue
            order += 1
            block_no += 1
            table_no += 1
            table_id = f"T-{table_no:03d}"
            headers = [normalize_header(cell, idx) for idx, cell in enumerate(rows[0], 1)]
            parsed_rows = []
            for idx, row in enumerate(rows[1:], 1):
                item = {"row_id": f"{table_id}-R{idx:03d}", "cells": {}}
                for col, header in enumerate(headers):
                    item["cells"][header] = row[col] if col < len(row) else ""
                parsed_rows.append(item)
            table = {
                "table_id": table_id,
                "table_type_guess": guess_table_type(headers, section_path),
                "section_path": list(section_path),
                "headers": headers,
                "rows": parsed_rows,
            }
            tables.append(table)
            blocks.append(
                {
                    "block_id": f"B-{block_no:04d}",
                    "block_type": "table",
                    "table_id": table_id,
                    "table_type_guess": table["table_type_guess"],
                    "section_path": list(section_path),
                    "order": order,
                    "headers": headers,
                    "rows": parsed_rows,
                }
            )
    image_manifest = extract_all_images(src, images_dir, blocks)
    return blocks, tables, image_manifest


def parse_docx_with_xml(src: Path, images_dir: Path | None = None) -> tuple[list[dict], list[dict], list[dict]]:
    with zipfile.ZipFile(src) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
        rels = parse_relationships(zf)
    blocks: list[dict] = []
    tables: list[dict] = []
    section_path: list[str] = []
    block_no = 0
    table_no = 0
    order = 0
    image_no = 0
    image_manifest: list[dict] = []

    body = root.find("w:body", NS)
    if body is None:
        return blocks, tables, image_manifest

    for child in list(body):
        tag = strip_ns(child.tag)
        if tag == "p":
            text = text_of(child)
            if not text:
                continue
            order += 1
            block_no += 1
            block_id = f"B-{block_no:04d}"
            level = heading_level(child)
            if level:
                section_path = section_path[:level - 1] + [text]
                block_type = "heading"
                block = {
                    "block_id": block_id,
                    "block_type": block_type,
                    "level": level,
                    "text": text,
                    "section_path": list(section_path),
                    "order": order,
                }
            else:
                block_type = "paragraph"
                block = {
                    "block_id": block_id,
                    "block_type": block_type,
                    "text": text,
                    "section_path": list(section_path),
                    "order": order,
                }
            image_refs = parse_block_images(child, rels)
            if image_refs:
                manifest = append_images_for_refs(
                    images_dir,
                    image_refs,
                    block_id=block_id,
                    table_id=None,
                    section_path=section_path,
                    order=order,
                    image_no_ref=[image_no],
                    source_doc=src,
                )
                if manifest:
                    image_no += len(manifest)
                    image_manifest.extend(manifest)
                    block["image_ids"] = [item["image_id"] for item in manifest]
            blocks.append(block)
        elif tag == "tbl":
            order += 1
            block_no += 1
            table_no += 1
            table_id = f"T-{table_no:03d}"
            headers, rows = parse_table(child, table_id)
            table = {
                "table_id": table_id,
                "table_type_guess": guess_table_type(headers, section_path),
                "section_path": list(section_path),
                "headers": headers,
                "rows": rows,
            }
            tables.append(table)
            block = {
                "block_id": f"B-{block_no:04d}",
                "block_type": "table",
                "table_id": table_id,
                "table_type_guess": table["table_type_guess"],
                "section_path": list(section_path),
                "order": order,
                "headers": headers,
                "rows": rows,
            }
            image_refs = parse_block_images(child, rels)
            if image_refs:
                manifest = append_images_for_refs(
                    images_dir,
                    image_refs,
                    block_id=block["block_id"],
                    table_id=table_id,
                    section_path=section_path,
                    order=order,
                    image_no_ref=[image_no],
                    source_doc=src,
                )
                if manifest:
                    image_no += len(manifest)
                    image_manifest.extend(manifest)
                    block["image_ids"] = [item["image_id"] for item in manifest]
            blocks.append(block)
    return blocks, tables, image_manifest


def iter_doc_blocks(doc: DocumentObject):
    body = doc.element.body
    paragraph_map = {paragraph._p: paragraph for paragraph in doc.paragraphs}
    table_map = {table._tbl: table for table in doc.tables}
    for child in body.iterchildren():
        if child in paragraph_map:
            yield paragraph_map[child]
        elif child in table_map:
            yield table_map[child]


def paragraph_heading_level(paragraph: Paragraph) -> int | None:
    style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
    match = re.search(r"Heading\s*(\d+)|标题\s*(\d+)", style_name, re.I)
    if match:
        return int(match.group(1) or match.group(2))
    text = (paragraph.text or "").strip()
    if re.match(r"^\d+(?:\.\d+){0,4}\s+", text):
        return min(text.split(" ", 1)[0].count(".") + 1, 6)
    return None


def extract_all_images(src: Path, images_dir: Path | None, blocks: list[dict]) -> list[dict]:
    if images_dir is None:
        return []
    entries: list[dict] = []
    with zipfile.ZipFile(src) as zf:
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
        for idx, name in enumerate(media, 1):
            ext = Path(name).suffix.lower() or ".bin"
            image_id = f"IMG-{idx:03d}"
            filename = f"{image_id}{ext}"
            (images_dir / filename).write_bytes(zf.read(name))
            near = blocks[-1] if blocks else {}
            entries.append(
                {
                    "image_id": image_id,
                    "order": near.get("order", idx),
                    "section_path": near.get("section_path", []),
                    "near_block_id": near.get("block_id"),
                    "near_table_id": near.get("table_id"),
                    "source": name,
                    "filename": f"parsed/images/{filename}",
                    "extension": ext.lstrip("."),
                }
            )
    return entries


def parse_relationships(zf: zipfile.ZipFile) -> dict[str, str]:
    try:
        rels_root = ET.fromstring(zf.read("word/_rels/document.xml.rels"))
    except KeyError:
        return {}
    mapping: dict[str, str] = {}
    for rel in rels_root.findall(f".//{{{REL_NS}}}Relationship"):
        rel_id = rel.attrib.get("Id")
        target = rel.attrib.get("Target")
        if rel_id and target:
            mapping[rel_id] = target
    return mapping


def parse_block_images(node: ET.Element, rels: dict[str, str]) -> list[str]:
    refs: list[str] = []
    for blip in node.findall(".//a:blip", NS_EXT):
        rel_id = blip.attrib.get(f"{{{NS_EXT['r']}}}embed")
        if rel_id and rel_id in rels:
            refs.append(rels[rel_id])
    return list(dict.fromkeys(refs))


def append_images_for_refs(
    images_dir: Path | None,
    refs: list[str],
    *,
    block_id: str,
    table_id: str | None,
    section_path: list[str],
    order: int,
    image_no_ref: list[int],
    source_doc: Path,
) -> list[dict]:
    if images_dir is None:
        return []
    entries: list[dict] = []
    for ref in refs:
        image_no_ref[0] += 1
        image_id = f"IMG-{image_no_ref[0]:03d}"
        rel_path = f"word/{ref}" if not ref.startswith("word/") else ref
        ext = Path(ref).suffix.lower() or ".bin"
        filename = f"{image_id}{ext}"
        if maybe_extract_image(source_doc, rel_path, images_dir / filename):
            entries.append(
                {
                    "image_id": image_id,
                    "order": order,
                    "section_path": list(section_path),
                    "near_block_id": block_id,
                    "near_table_id": table_id,
                    "source": rel_path,
                    "filename": f"parsed/images/{filename}",
                    "extension": ext.lstrip("."),
                }
            )
    return entries


def maybe_extract_image(src: Path, rel_path: str, dest: Path) -> bool:
    with zipfile.ZipFile(src) as zf:
        try:
            data = zf.read(rel_path)
        except KeyError:
            return False
    dest.write_bytes(data)
    return True


def parse_table(tbl: ET.Element, table_id: str) -> tuple[list[str], list[dict]]:
    raw_rows: list[list[str]] = []
    for tr in tbl.findall("w:tr", NS):
        cells = [text_of(tc) for tc in tr.findall("w:tc", NS)]
        if any(cells):
            raw_rows.append(cells)
    if not raw_rows:
        return [], []
    headers = [normalize_header(c, idx) for idx, c in enumerate(raw_rows[0], 1)]
    rows: list[dict] = []
    for idx, row in enumerate(raw_rows[1:], 1):
        item = {"row_id": f"{table_id}-R{idx:03d}", "cells": {}}
        for col, header in enumerate(headers):
            item["cells"][header] = row[col] if col < len(row) else ""
        rows.append(item)
    return headers, rows


def text_of(node: ET.Element) -> str:
    parts = [t.text or "" for t in node.findall(".//w:t", NS)]
    return re.sub(r"\s+", " ", "".join(parts)).strip()


def heading_level(p: ET.Element) -> int | None:
    style = p.find(".//w:pStyle", NS)
    if style is None:
        return None
    val = style.attrib.get(f"{{{NS['w']}}}val", "")
    match = re.search(r"Heading(\d+)|heading\s*(\d+)", val, re.I)
    if not match:
        return None
    return int(match.group(1) or match.group(2))


def normalize_header(value: str, idx: int) -> str:
    value = value.strip() or f"Column{idx}"
    return re.sub(r"\s+", "_", value)


def guess_table_type(headers: list[str], section_path: list[str]) -> str:
    text = " ".join(headers + section_path)
    if any(k in text for k in ["权限", "权限", "数据边界"]):
        return "permission_table"
    if any(k in text for k in ["字段", "类型", "属性"]):
        return "field_table"
    if any(k in text for k in ["功能", "模块", "按钮"]):
        return "function_table"
    return "general_table"


def write_markdown(path: Path, blocks: list[dict]) -> None:
    lines: list[str] = []
    for block in blocks:
        if block["block_type"] == "heading":
            lines.append("#" * int(block["level"]) + f" [{block['block_id']}] " + block["text"])
        elif block["block_type"] == "paragraph":
            lines.append(f"[{block['block_id']}] {block['text']}")
        elif block["block_type"] == "table":
            lines.append(f"[{block['table_id']}][{block['block_id']}] {', '.join(block.get('headers', []))}")
            headers = ["row_id"] + list(block.get("headers", []))
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join("---" for _ in headers) + " |")
            for row in block.get("rows", []):
                cells = row.get("cells", {})
                values = [row.get("row_id", "")] + [str(cells.get(h, "")) for h in block.get("headers", [])]
                lines.append("| " + " | ".join(v.replace("|", "\\|") for v in values) + " |")
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_yaml(path: Path, data: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(yaml.safe_dump(data, allow_unicode=True, sort_keys=False), encoding="utf-8")


def strip_ns(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def detect_parser_warnings(src: Path) -> list[dict[str, str]]:
    warnings: list[dict[str, str]] = []
    with zipfile.ZipFile(src) as zf:
        names = zf.namelist()
        if not any(name.startswith("word/media/") for name in names):
            warnings.append({"warning_type": "image_not_found", "message": "No embedded image references found in document."})
        else:
            warnings.append({"warning_type": "image_parsed", "message": "Image relationships discovered and extracted when possible."})
        if "word/comments.xml" in names:
            warnings.append({"warning_type": "comments_not_parsed", "message": "Docx comments are not parsed in current docx stage."})
        if any(name.startswith("word/header") or name.startswith("word/footer") for name in names):
            warnings.append({"warning_type": "header_footer_not_parsed", "message": "Document headers/footers are intentionally not parsed."})
    return warnings


if __name__ == "__main__":
    raise SystemExit(main())
